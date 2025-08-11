import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import os

PDF_PATH = '/home/gyasis/Desktop/Data_resume.pdf'
OUTPUT_DOCX = 'Data_resume_two_column.docx'

# Helper to add two-column section to docx
def add_two_column_section(doc):
    sectPr = doc.sections[-1]._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

# Extract text from PDF, attempting to preserve two columns
# This is a best-effort approach; perfect column detection is not always possible

def extract_columns_text(pdf_path):
    text_col1 = []
    text_col2 = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            # Try to split the page into two columns
            width = page.width
            mid = width / 2
            # Left column
            left = page.within_bbox((0, 0, mid, page.height))
            # Right column
            right = page.within_bbox((mid, 0, width, page.height))
            col1_text = left.extract_text() if left else ''
            col2_text = right.extract_text() if right else ''
            text_col1.append(col1_text or '')
            text_col2.append(col2_text or '')
    return text_col1, text_col2

def main():
    if not os.path.exists(PDF_PATH):
        print(f"PDF not found: {PDF_PATH}")
        return
    print("Extracting text from PDF...")
    col1, col2 = extract_columns_text(PDF_PATH)
    print("Creating Word document...")
    doc = Document()
    add_two_column_section(doc)
    # Insert text: left column, then right column (duplicated if needed)
    for i in range(len(col1)):
        p1 = doc.add_paragraph(col1[i])
        p1.style.font.size = Pt(10)
        p2 = doc.add_paragraph(col2[i])
        p2.style.font.size = Pt(10)
    doc.save(OUTPUT_DOCX)
    print(f"Saved as {OUTPUT_DOCX}")

if __name__ == '__main__':
    main() 